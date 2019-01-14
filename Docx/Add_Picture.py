import docx

doc = docx.Document()
doc.add_paragraph("This is it!!!")
doc.add_picture("Files\\zophie.png",height=docx.shared.Inches(1),width=docx.shared.Cm(4))
# width and height keyword arguments are Optional
doc.add_paragraph("Apna time Aega!!!")
doc.save("Files\\PicturedDoc.docx")