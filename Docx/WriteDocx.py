import docx

doc = docx.Document()
doc.add_paragraph("Hello World !!!")

""" Both add_paragraph() and add_run() accept an optional second argument that is a string of the Paragraph
or Run object’s style """

paraObj1 = doc.add_paragraph("Second Paragraph. ","Title")
paraObj1.add_run("Appended with Second Paragraph")
doc.paragraphs[1].runs[0].add_break()   # add_break(docx.enum.text.WD_BREAK.PAGE) for PAGE BREAK. Method to runs Object
paraObj2 = doc.add_paragraph("Next Page. Third Paragraph")

doc.add_heading('Header 0', 0)      # 0 is the biggest. 0 to 4
doc.add_heading('Header 2', 2)      # 2 is bigger than 4. Color is Blue.
doc.add_heading('Header 4', 4)      # Heading is placed wherever it is placed in the code.

doc.save("Files\\WriteDocx.docx")