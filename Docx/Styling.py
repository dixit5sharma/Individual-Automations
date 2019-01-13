#! python3
import docx

doc = docx.Document("Files\\demo.docx")
print(doc.paragraphs[0].text)
doc.paragraphs[0].style = "Normal"
print(doc.paragraphs[0].style)
print(doc.paragraphs[1].runs[0].text)
doc.paragraphs[1].runs[0].style = "Quote Char"
doc.paragraphs[1].runs[1].underline = True
doc.paragraphs[1].runs[3].underline = True

doc.save("Files\\ReStyled.docx")

"""" For Word documents, there are three types of styles: Paragraph styles can be applied to Paragraph objects,
character styles can be applied to Run objects, and linked styles can be applied to both kinds of objects. You
can give both Paragraph and Run objects styles by setting their style attribute to a string. This string should be
the name of a style. If style is set to None, then there will be no style associated with the Paragraph or Run object.
The string values for the default Word styles are as follows:
'Normal'    'BodyText'  'BodyText2' 'BodyText3' 'Caption'   'Heading1'  'Heading2'  'Heading3'  'Heading4'
'Heading5'  'Heading6'  'Heading7'  'Heading8'  'Heading9'  'IntenseQuote'  'List'  'List2' 'List3' 'ListBullet'
'ListBullet2'   'ListBullet3'   'ListContinue'  'ListContinue2' 'ListContinue3' 'ListNumber'    'ListNumber2'
'ListNumber3'   'ListParagraph' 'MacroText' 'NoSpacing' 'Quote' 'Subtitle'  'TOCHeading'    'Title'  """


"""" Runs can be further styled using text attributes. Each attribute can be set to one of three values: True 
(the attribute is always enabled, no matter what other styles are applied to the run), False (the attribute is 
always disabled), or None (defaults to whatever the run’s style is set to). 
bold            The text appears in bold.
italic          The text appears in italic.
underline       The text is underlined.
strike          The text appears with strikethrough.
double_strike   The text appears with double strikethrough.
all_caps        The text appears in capital letters.
small_caps      The text appears in capital letters, with lowercase letters two points smaller.
shadow          The text appears with a shadow.
outline         The text appears outlined rather than solid.
rtl             The text is written right-to-left.
imprint         The text appears pressed into the page.
emboss          The text appears raised off the page in relief 
"""