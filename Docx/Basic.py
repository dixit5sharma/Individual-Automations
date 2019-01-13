#! python3
""" The Document object contains a list of Paragraph objects for the paragraphs (Enter based) in the document.
Each of these Paragraph objects contains a list of one or more Run objects. A Run object is a contiguous run
of text with the same style. A new Run object is needed whenever the text style changes """
import docx

doc = docx.Document("Files\\demo.docx")     # <class 'docx.document.Document'>
print(len(doc.paragraphs))      # int

# def mapThis(a):
#     return a.text
# b = list(map(mapThis,doc.paragraphs))     # NICEEEEEEEE - Applying a function/method to each element of the list

b = [item.text for item in doc.paragraphs]  # ULTIMATE - EVEN BETTER - LIST COMPREHENSION IS THE BEST
print(b)         # Last one is empty and is the 7th one

print(doc.paragraphs[1].text)
print(doc.paragraphs[1].runs)       # list of <docx.text.run.Run> objects

print(len(doc.paragraphs[1].runs))
print(doc.paragraphs[1].runs[0].text)       # Run is identified as a change in style within a paragraph.
c = [item.text for item in doc.paragraphs[1].runs]      # LIST COMPREHENTION again
print(c)